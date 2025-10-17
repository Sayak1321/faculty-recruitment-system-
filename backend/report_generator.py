
import os
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from backend import db
import json

def generate_pdf_report(job_id):
    job = db.get_job(job_id)
    apps = db.get_applications_by_job(job_id)
    if not job:
        raise ValueError("Job not found")
    filename = f"report_job_{job_id}.pdf"
    filepath = os.path.join(db.REPORTS_DIR, filename)

    doc = SimpleDocTemplate(filepath, pagesize=A4, rightMargin=30,leftMargin=30, topMargin=30,bottomMargin=18)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(f"Recruitment Report: {job['title']}", styles['Title']))
    story.append(Spacer(1,12))
    story.append(Paragraph(f"Department: {job['department']}", styles['Normal']))
    story.append(Spacer(1,12))
    # table header
    data = [["Candidate", "Email", "Score", "Eligible", "Status"]]
    for a in apps:
        parsed = json.loads(a['parsed_json']) if a['parsed_json'] else {}
        data.append([a['candidate_name'], a['email'] or "", str(a['score'] or "-"), "Yes" if a['eligible'] else "No", a['status'] or ""])
    table = Table(data, colWidths=[120,130,60,60,100])
    table.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.grey),
                               ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
                               ('ALIGN',(0,0),(-1,-1),'LEFT'),
                               ('GRID',(0,0),(-1,-1),0.5,colors.black)]))
    story.append(table)
    story.append(Spacer(1,12))
    story.append(Paragraph("Detailed evaluations available in the system.", styles['Normal']))
    doc.build(story)
    # save record
    db.insert_report(job_id, filepath)
    return filepath

def generate_docx_report(job_id):
    job = db.get_job(job_id)
    apps = db.get_applications_by_job(job_id)
    if not job:
        raise ValueError("Job not found")
    filename = f"report_job_{job_id}.docx"
    filepath = os.path.join(db.REPORTS_DIR, filename)
    doc = Document()
    doc.add_heading(f"Recruitment Report: {job['title']}", level=1)
    doc.add_paragraph(f"Department: {job['department']}")
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Candidate'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Score'
    hdr_cells[3].text = 'Eligible'
    hdr_cells[4].text = 'Status'
    for a in apps:
        row_cells = table.add_row().cells
        row_cells[0].text = a['candidate_name']
        row_cells[1].text = a['email'] or ""
        row_cells[2].text = str(a['score'] or "-")
        row_cells[3].text = "Yes" if a['eligible'] else "No"
        row_cells[4].text = a['status'] or ""
    doc.add_paragraph("\nDetailed evaluations are in the system database.")
    doc.save(filepath)
    db.insert_report(job_id, filepath)
    return filepath
